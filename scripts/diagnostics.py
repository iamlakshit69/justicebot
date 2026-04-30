#!/usr/bin/env python3
"""
JusticeBot Diagnostics Runner (Dark Theme, Comprehensive)
Generates dark-themed PDF, DOCX, and JSON reports with multiple charts.
Tests all application modules, utilities, agents, and endpoints.
"""

import os
import sys
import json
import time
import uuid
import traceback
import importlib
import datetime
from pathlib import Path

# Analytics Libs
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage

# ── Global State ──────────────────────────────────────────────────────────────
results = []
metrics = {"latencies": {}}
start_ts = datetime.datetime.now()
report_dir = Path("diagnostics") / start_ts.strftime("%Y-%m-%d_%H-%M-%S")
report_dir.mkdir(parents=True, exist_ok=True)

def record(category: str, test: str, status: str, detail: str = "", duration_ms: float = 0):
    results.append({
        "category": category, 
        "test": test, 
        "status": status, 
        "detail": detail,
        "duration_ms": duration_ms
    })
    
    # Store latency for route/agent charts if duration > 0
    if duration_ms > 0 and (category == "Routes" or category == "Agents"):
        metrics["latencies"][test] = duration_ms
        
    color_map = {"PASS": "\033[32m", "FAIL": "\033[31m", "WARN": "\033[33m", "INFO": "\033[36m"}
    reset = "\033[0m"
    print(f"{color_map.get(status, '')}[{status:4s}]{reset} {category}: {test}")

# ── Tests ───────────────────────────────────────────────────────────────────
def test_environment_and_config():
    record("Env", f"Python {sys.version_info.major}.{sys.version_info.minor}", "PASS" if sys.version_info >= (3,9) else "WARN")
    
    deps = ["flask", "dotenv", "groq", "fitz", "docx", "fpdf", "matplotlib", "reportlab"]
    for pkg in deps:
        try:
            importlib.import_module(pkg)
            record("Env", f"Dependency {pkg} loaded", "PASS")
        except ImportError:
            record("Env", f"Dependency {pkg} missing", "FAIL")

    try:
        import config as cfg
        record("Config", "config.py loaded", "PASS")
        req_attrs = ["GROQ_API_KEY", "MODELS", "TEMPERATURE", "APP_TITLE"]
        for attr in req_attrs:
            if hasattr(cfg, attr): record("Config", f"{attr} present", "PASS")
            else: record("Config", f"Missing {attr} in config.py", "FAIL")
    except Exception as e:
        record("Config", "Failed to load config", "FAIL", str(e))

def test_modules():
    core_modules = [
        "memory.session", "data.dlsa_offices", "utils.dlsa", 
        "utils.doc_parser", "utils.pdf_export", "agents.router", 
        "agents.legal_advisor", "agents.doc_analyzer", "agents.drafter",
        "prompts.legal_prompts", "prompts.draft_prompts"
    ]
    for mod in core_modules:
        try:
            importlib.import_module(mod)
            record("Modules", f"Module {mod} built correctly", "PASS")
        except Exception as e:
            record("Modules", f"Module {mod} failed", "FAIL", str(e))

def test_utilities():
    # Session
    try:
        from memory.session import create_session, get_session
        sess = create_session("test_sess")
        record("Utils", "Session created successfully", "PASS" if isinstance(sess, dict) else "FAIL")
    except Exception as e:
        record("Utils", "Session memory failed", "FAIL", str(e))

    # DLSA lookup
    try:
        from utils.dlsa import get_dlsa
        dlsa = get_dlsa("delhi")
        record("Utils", "DLSA Lookup works (Delhi)", "PASS" if dlsa else "FAIL")
    except Exception as e:
        record("Utils", "DLSA Lookup failed", "FAIL", str(e))

def test_routes():
    try:
        import app as app_module
        app = app_module.app
        app.config["TESTING"] = True
        
        with app.test_client() as client:
            t0 = time.time()
            resp = client.get("/")
            t1 = time.time()
            record("Routes", "GET / returns 200 OK", "PASS" if resp.status_code == 200 else "FAIL", duration_ms=round((t1-t0)*1000, 2))
             
            t0 = time.time()
            resp_dlsa = client.get("/api/dlsa?state=punjab")
            t1 = time.time()
            record("Routes", "GET /api/dlsa valid for Punjab", "PASS" if resp_dlsa.status_code == 200 else "FAIL", duration_ms=round((t1-t0)*1000, 2))
            
            t0 = time.time()
            resp_fail_dlsa = client.get("/api/dlsa?state=unknown_state")
            t1 = time.time()
            record("Routes", "GET /api/dlsa handles 404 cleanly", "PASS" if resp_fail_dlsa.status_code == 404 else "FAIL", duration_ms=round((t1-t0)*1000, 2))

            t0 = time.time()
            resp_draft = client.post("/api/draft", json={"draft_type": "notice", "session_id": "test"})
            t1 = time.time()
            # It will fail at 400 since no valid session memory exists with valid context, which is expected behaviour.
            record("Routes", "POST /api/draft handles missing context properly", "PASS" if resp_draft.status_code == 400 else "FAIL", duration_ms=round((t1-t0)*1000, 2))
            
    except Exception as e:
         record("Routes", "App routing tests crashed", "FAIL", str(e))

# ── Generators ──────────────────────────────────────────────────────────────

def generate_charts():
    # Set dark theme globally for plots
    plt.style.use('dark_background')
    
    # 1. Overall Status Bar Chart
    chart1_path = report_dir / "chart_overview.png"
    counts = {"PASS": 0, "FAIL": 0, "WARN": 0, "INFO": 0}
    for r in results: counts[r["status"]] += 1
    
    labels = list(counts.keys())
    values = list(counts.values())
    colors_list = ['#00E676', '#FF1744', '#FFEA00', '#00B0FF'] # Bright neon for dark theme
    
    fig, ax = plt.subplots(figsize=(7, 4))
    fig.patch.set_facecolor('#1E1E1E')
    ax.set_facecolor('#1E1E1E')
    bars = ax.bar(labels, values, color=colors_list)
    ax.set_title("Diagnostics Overview Analysis", color='white', fontweight='bold')
    ax.set_ylabel("Quantity", color='#CCCCCC')
    
    for bar in bars:
        height = bar.get_height()
        ax.annotate(f'{int(height)}',
                    xy=(bar.get_x() + bar.get_width() / 2, height),
                    xytext=(0, 3), textcoords="offset points",
                    ha='center', va='bottom', color='white', fontweight='bold')
    
    plt.tight_layout()
    plt.savefig(chart1_path, dpi=300, facecolor='#1E1E1E')
    plt.close()
    
    # 2. Performance / Latency Chart
    chart2_path = report_dir / "chart_performance.png"
    latencies = metrics.get("latencies", {})
    if latencies:
        fig, ax = plt.subplots(figsize=(7, 4))
        fig.patch.set_facecolor('#1E1E1E')
        ax.set_facecolor('#1E1E1E')
        
        route_names = [r.replace("GET ", "").replace("POST ", "").split(" ")[0][:20] for r in latencies.keys()]
        times = list(latencies.values())
        
        y_pos = np.arange(len(route_names))
        ax.barh(y_pos, times, color='#D500F9')
        ax.set_yticks(y_pos)
        ax.set_yticklabels(route_names, color='#CCCCCC')
        ax.set_xlabel("Latency (ms)", color='#CCCCCC')
        ax.set_title("Test Latency Analysis", color='white', fontweight='bold')
        
        for i, v in enumerate(times):
            ax.text(v + 1, i, f"{v}ms", color='white', va='center', fontweight='bold')

        plt.tight_layout()
        plt.savefig(chart2_path, dpi=300, facecolor='#1E1E1E')
        plt.close()
    
    return str(chart1_path), (str(chart2_path) if latencies else None)

def generate_pdf(chart1_path, chart2_path):
    pdf_path = report_dir / "report.pdf"
    
    # Dark Theme Colors
    bg_color = colors.HexColor('#121212')
    text_color = colors.HexColor('#E0E0E0')
    accent_color = colors.HexColor('#00E676')
    
    def add_dark_bg(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(bg_color)
        canvas.rect(0, 0, letter[0], letter[1], fill=True, stroke=False)
        canvas.restoreState()

    doc = SimpleDocTemplate(str(pdf_path), pagesize=letter)
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle('TitleStyle', parent=styles['Heading1'], alignment=1, textColor=accent_color)
    subtitle_style = ParagraphStyle('SubStyle', parent=styles['Normal'], alignment=1, textColor=colors.HexColor('#AAAAAA'))
    section_style = ParagraphStyle('SectionStyle', parent=styles['Heading2'], textColor=colors.HexColor('#00B0FF'))
    
    elements = []
    
    elements.append(Paragraph("JusticeBot Total Infrastructure Report", title_style))
    elements.append(Spacer(1, 10))
    elements.append(Paragraph(f"Execution Timestamp: {start_ts.strftime('%Y-%m-%d %H:%M:%S')}", subtitle_style))
    elements.append(Spacer(1, 20))
    
    if os.path.exists(chart1_path):
        elements.append(RLImage(chart1_path, width=420, height=240))
        elements.append(Spacer(1, 10))
        
    if chart2_path and os.path.exists(chart2_path):
        elements.append(RLImage(chart2_path, width=420, height=240))
        elements.append(Spacer(1, 20))
    
    elements.append(Paragraph("Diagnostic Granular Log", section_style))
    
    data = [["Category", "Test Scenario", "Result"]]
    for r in results:
        status_text = r["status"]
        if status_text == "PASS": status_text = "PASS"
        data.append([r["category"], r["test"][:60], r["status"]])
        
    t = Table(data, colWidths=[100, 320, 60])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1E1E1E')),
        ('TEXTCOLOR', (0,0), (-1,0), accent_color),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0,0), (-1,0), 10),
        
        ('BACKGROUND', (0,1), (-1,-1), colors.HexColor('#222222')),
        ('TEXTCOLOR', (0,1), (-1,-1), text_color),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#333333')),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    
    elements.append(t)
    doc.build(elements, onFirstPage=add_dark_bg, onLaterPages=add_dark_bg)

def generate_docx(chart1_path, chart2_path):
    docx_path = report_dir / "report.docx"
    doc = Document()
    
    # Basic Word dark styling doesn't apply cleanly without a template file,
    # but we can set heading and text colors
    head = doc.add_heading('JusticeBot Infrastructure Metrics', 0)
    head.alignment = 1
    doc.add_paragraph(f"Timestamp: {start_ts.strftime('%Y-%m-%d %H:%M:%S')}")
    
    if os.path.exists(chart1_path):
         doc.add_picture(chart1_path, width=Inches(5.0))
    if chart2_path and os.path.exists(chart2_path):
         doc.add_picture(chart2_path, width=Inches(5.0))
        
    doc.add_heading('Test Execution Log', level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Dark List Accent 1' # A built-in dark word style
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Module"
    hdr_cells[1].text = "Test Executed"
    hdr_cells[2].text = "State"
    
    for r in results:
        row = table.add_row().cells
        row[0].text = r["category"]
        row[1].text = r["test"]
        row[2].text = r["status"]
        
    doc.save(docx_path)

def generate_json():
    json_path = report_dir / "report.json"
    counts = {"PASS": 0, "FAIL": 0, "WARN": 0, "INFO": 0}
    for r in results: counts[r["status"]] += 1
    
    data = {
        "metadata": {
            "application": "JusticeBot",
            "generated_at": start_ts.isoformat(),
        }, 
        "metrics": metrics,
        "summary": counts, 
        "details": results
    }
    with open(json_path, 'w') as f: 
        json.dump(data, f, indent=2)

if __name__ == "__main__":
    print(f"=======================================")
    print(f"JusticeBot Supreme Diagnostics Engine")
    print(f"=======================================\n")
    
    test_environment_and_config()
    test_modules()
    test_utilities()
    test_routes()
    
    chart1, chart2 = generate_charts()
    generate_pdf(chart1, chart2)
    generate_docx(chart1, chart2)
    generate_json()
    
    print(f"\n=======================================")
    print(f"Outputs generated natively at:")
    print(f" -> {report_dir}")
    print(f"=======================================")
